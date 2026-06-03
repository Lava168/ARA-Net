#!/usr/bin/env python3
"""Audit a Word manuscript for stale ARA-Net V6 claim-boundary language.

The public Markdown package already has a claim-boundary audit. This script is
for the submitted Word manuscript, where old CAS, attention, Braak, OASIS, and
clinical-deployment wording can survive after the repository has been updated.
"""
from __future__ import annotations

import argparse
import re
from dataclasses import dataclass
from pathlib import Path

from docx import Document


DEFAULT_DOCX = Path("/Users/mac/Downloads/ARA-Net_MedIA_Paper（完整版）.docx")
DEFAULT_OUTPUT = Path("reports/v6_final_model/word_manuscript_claim_audit.md")


SAFE_CUES = [
    "not",
    "no ",
    "no longer",
    "do not",
    "does not",
    "did not",
    "cannot",
    "avoid",
    "removed",
    "remove",
    "replaced",
    "replace",
    "rather than",
    "instead",
    "limitation",
    "stress test",
    "stress-test",
    "unresolved",
    "weak",
    "failed",
    "failure",
    "below",
    "non-significant",
    "unsupported",
    "insufficient",
    "proxy",
    "not direct",
    "not intended",
    "not a medical device",
    "not cleared",
    "not approved",
    "requires",
    "required",
    "research prototype",
]


@dataclass(frozen=True)
class Rule:
    name: str
    severity: str
    pattern: re.Pattern[str]
    recommendation: str
    allow_safe_context: bool = True


RULES = [
    Rule(
        "attention_cas_biomarker_claim",
        "blocker",
        re.compile(
            r"\b(CAS|Clinical Alignment Score|Attention-as-Biomarker|attention)[^.]{0,180}\b"
            r"(validat|biomarker|clinical alignment|Braak|neuropatholog|disease progression|diagnostic explanation|interpretability)",
            re.I,
        ),
        "Replace attention/CAS biomarker language with atlas structural neurodegeneration consistency.",
        False,
    ),
    Rule(
        "direct_braak_claim",
        "blocker",
        re.compile(r"\b(Braak|neuropatholog)[^.]{0,180}\b(validat|correlat|proof|support|consistent|align|correspond)", re.I),
        "Remove direct Braak-stage validation language; use MRI neurodegeneration proxy wording.",
        False,
    ),
    Rule(
        "oasis_success_claim",
        "blocker",
        re.compile(r"\bOASIS[^.]{0,180}\b(success|successful|validat|strong|generaliz|robust|interpretability)", re.I),
        "Report OASIS only as an unresolved stress-test limitation.",
    ),
    Rule(
        "zero_shot_success_claim",
        "blocker",
        re.compile(r"\b(pure\s+)?zero[- ]shot[^.]{0,120}\b(success|successful|validat|strong|generaliz|robust)", re.I),
        "Do not claim pure zero-shot success; use domain-adapted external heldout wording.",
    ),
    Rule(
        "clinical_ready_claim",
        "blocker",
        re.compile(
            r"\b(deployment[- ]ready|clinically deployable|clinical deployment|clinical device|diagnostic device|medical device|clinical adoption|clinically desirable|clinical trust|AI-assisted AD diagnosis)",
            re.I,
        ),
        "Frame as an open-source research prototype, not a deployment-ready clinical device.",
    ),
    Rule(
        "old_attention_model_framing",
        "blocker",
        re.compile(r"\b(Atlas-guided Region Attention Network|attention-guided ARA-Net|multi-head self-attention|atlas-guided region attention|attention weights|attention maps)\b", re.I),
        "Reframe the manuscript around the final atlas-guided multimodal subject-level rescue ensemble.",
        False,
    ),
    Rule(
        "old_primary_dataset_protocol",
        "warning",
        re.compile(r"\b(2,401|six-seed|five-fold|30 runs|1,500 test samples|887 unlabeled MRI volumes)\b", re.I),
        "Check whether this old v3 protocol text should be replaced by V6 subject-level ADNI/AIBL/IXI/OASIS protocol text.",
    ),
    Rule(
        "reference_placeholder",
        "blocker",
        re.compile(r"Error!\s+Reference source not found|REF _Ref|Equation\s+Error", re.I),
        "Fix broken Word references, stale equation fields, or unresolved cross-references.",
    ),
]


def paragraph_texts(docx_path: Path) -> list[str]:
    document = Document(str(docx_path))
    texts: list[str] = []
    for paragraph in document.paragraphs:
        text = " ".join(paragraph.text.split())
        if text:
            texts.append(text)
    for table in document.tables:
        for row in table.rows:
            cells = [" ".join(cell.text.split()) for cell in row.cells]
            text = " | ".join(cell for cell in cells if cell)
            if text:
                texts.append(text)
    return texts


def context_is_safe(texts: list[str], index: int) -> bool:
    start = max(0, index - 2)
    end = min(len(texts), index + 3)
    context = " ".join(texts[start:end]).lower()
    return any(cue in context for cue in SAFE_CUES)


def audit_texts(texts: list[str]) -> list[dict]:
    findings: list[dict] = []
    for index, text in enumerate(texts):
        for rule in RULES:
            if rule.pattern.search(text):
                safe = rule.allow_safe_context and context_is_safe(texts, index)
                status = "allowed_safe_context" if safe else "flagged"
                findings.append(
                    {
                        "index": index + 1,
                        "rule": rule.name,
                        "severity": rule.severity,
                        "status": status,
                        "text": text,
                        "recommendation": rule.recommendation,
                    }
                )
    return findings


def truncate(text: str, limit: int = 260) -> str:
    return text if len(text) <= limit else text[: limit - 3] + "..."


def build_report(docx_path: Path, texts: list[str], findings: list[dict]) -> str:
    blockers = [item for item in findings if item["status"] == "flagged" and item["severity"] == "blocker"]
    warnings = [item for item in findings if item["status"] == "flagged" and item["severity"] == "warning"]
    allowed = [item for item in findings if item["status"] == "allowed_safe_context"]
    status = "pass" if not blockers else "fail"
    lines = [
        "# Word Manuscript Claim Audit",
        "",
        f"- Source DOCX: `{docx_path}`",
        f"- Paragraph/table-text units scanned: {len(texts)}",
        f"- Status: **{status}**",
        f"- Blocker findings: {len(blockers)}",
        f"- Warning findings: {len(warnings)}",
        f"- Allowed safe-context mentions: {len(allowed)}",
        "",
        "## Interpretation",
        "",
        "This report audits the Word manuscript, not only the public repository Markdown. Blockers are stale or unsupported manuscript claims that should be removed before resubmission.",
        "",
    ]
    for title, items in [("Blockers", blockers), ("Warnings", warnings)]:
        lines += [f"## {title}", ""]
        if not items:
            lines.append(f"No {title.lower()} were detected.")
            lines.append("")
            continue
        lines += ["| unit | rule | text | recommended action |", "|---:|---|---|---|"]
        for item in items:
            lines.append(
                f"| {item['index']} | {item['rule']} | {truncate(item['text'])} | {item['recommendation']} |"
            )
        lines.append("")

    lines += [
        "## Allowed Safe-Context Examples",
        "",
        "| unit | rule | text |",
        "|---:|---|---|",
    ]
    for item in allowed[:40]:
        lines.append(f"| {item['index']} | {item['rule']} | {truncate(item['text'])} |")
    if len(allowed) > 40:
        lines.append(f"| ... | ... | {len(allowed) - 40} additional safe-context mentions omitted. |")
    lines.append("")
    return "\n".join(lines)


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--docx", type=Path, default=DEFAULT_DOCX)
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT)
    args = parser.parse_args()

    texts = paragraph_texts(args.docx)
    findings = audit_texts(texts)
    report = build_report(args.docx, texts, findings)
    args.output.parent.mkdir(parents=True, exist_ok=True)
    args.output.write_text(report, encoding="utf-8")
    print(f"[saved] {args.output}")
    blockers = [item for item in findings if item["status"] == "flagged" and item["severity"] == "blocker"]
    print(f"[blockers] {len(blockers)}")
    if blockers:
        raise SystemExit(2)


if __name__ == "__main__":
    main()
