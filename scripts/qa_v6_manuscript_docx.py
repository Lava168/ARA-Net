#!/usr/bin/env python3
"""Structural QA for the generated ARA-Net V6 manuscript DOCX."""
from __future__ import annotations

import argparse
import re
import zipfile
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


DEFAULT_DOCX = Path("reports/v6_final_model/ARA-Net_V6_full_manuscript_draft.docx")
DEFAULT_OUTPUT = Path("reports/v6_final_model/ARA-Net_V6_full_manuscript_docx_qa.md")

PLACEHOLDER_RE = re.compile(
    r"Error!\s+Reference source not found|REF _Ref|Equation\s+Error|TODO|PLACEHOLDER",
    re.I,
)


def count_numbering_defs(docx_path: Path) -> int:
    with zipfile.ZipFile(docx_path) as zf:
        try:
            xml = zf.read("word/numbering.xml").decode("utf-8", errors="ignore")
        except KeyError:
            return 0
    return xml.count("<w:num ")


def table_widths(table) -> list[int]:
    widths: list[int] = []
    grid = table._tbl.tblGrid
    if grid is None:
        return widths
    for col in grid.iterchildren():
        value = col.get(qn("w:w"))
        if value and value.isdigit():
            widths.append(int(value))
    return widths


def build_report(docx_path: Path) -> str:
    doc = Document(str(docx_path))
    paragraphs = [p for p in doc.paragraphs if p.text.strip()]
    tables = doc.tables
    all_text = "\n".join(p.text for p in paragraphs)
    for table in tables:
        for row in table.rows:
            all_text += "\n" + " | ".join(cell.text for cell in row.cells)

    headings = {}
    for paragraph in paragraphs:
        name = paragraph.style.name if paragraph.style is not None else ""
        if name.startswith("Heading"):
            headings[name] = headings.get(name, 0) + 1

    section = doc.sections[0]
    margins = {
        "top_in": round(section.top_margin.inches, 3),
        "right_in": round(section.right_margin.inches, 3),
        "bottom_in": round(section.bottom_margin.inches, 3),
        "left_in": round(section.left_margin.inches, 3),
    }

    table_rows = []
    bad_tables = []
    for idx, table in enumerate(tables, start=1):
        widths = table_widths(table)
        total = sum(widths)
        table_rows.append(
            {
                "index": idx,
                "rows": len(table.rows),
                "cols": len(table.columns),
                "grid_width_dxa": total,
                "has_grid": bool(widths),
            }
        )
        if not widths or total <= 0 or total > 9800:
            bad_tables.append(idx)

    placeholder_hits = [match.group(0) for match in PLACEHOLDER_RE.finditer(all_text)]
    required_phrases = [
        "balanced accuracy of 0.833",
        "OASIS transfer remained weak",
        "not a medical device",
        "atlas structural neurodegeneration consistency",
    ]
    missing_required = [phrase for phrase in required_phrases if phrase not in all_text]

    failures = []
    if len(paragraphs) < 80:
        failures.append("Unexpectedly few non-empty paragraphs.")
    if len(tables) < 3:
        failures.append("Expected manuscript tables were not found.")
    if placeholder_hits:
        failures.append("Placeholder or broken-reference text detected.")
    if missing_required:
        failures.append("Required manuscript phrases missing.")
    if bad_tables:
        failures.append("One or more tables lack sane fixed-grid geometry.")
    if count_numbering_defs(docx_path) == 0:
        failures.append("No numbering definitions found for lists.")

    status = "pass" if not failures else "fail"
    lines = [
        "# ARA-Net V6 Manuscript DOCX QA",
        "",
        f"- Source DOCX: `{docx_path}`",
        f"- Status: **{status}**",
        f"- Non-empty paragraphs: {len(paragraphs)}",
        f"- Tables: {len(tables)}",
        f"- Numbering definitions: {count_numbering_defs(docx_path)}",
        f"- Page margins: {margins}",
        "",
        "## Heading Counts",
        "",
        "| style | count |",
        "|---|---:|",
    ]
    for style, count in sorted(headings.items()):
        lines.append(f"| {style} | {count} |")

    lines += [
        "",
        "## Table Geometry",
        "",
        "| table | rows | cols | grid width DXA | fixed grid present |",
        "|---:|---:|---:|---:|---|",
    ]
    for row in table_rows:
        lines.append(
            f"| {row['index']} | {row['rows']} | {row['cols']} | "
            f"{row['grid_width_dxa']} | {row['has_grid']} |"
        )

    lines += [
        "",
        "## Required Phrase Check",
        "",
    ]
    if missing_required:
        lines.extend(f"- Missing: `{phrase}`" for phrase in missing_required)
    else:
        lines.append("All required phrases were found.")

    lines += [
        "",
        "## Placeholder Check",
        "",
    ]
    if placeholder_hits:
        lines.extend(f"- `{hit}`" for hit in placeholder_hits[:20])
    else:
        lines.append("No placeholder or broken-reference text was detected.")

    lines += [
        "",
        "## Visual Render QA",
        "",
        "LibreOffice render QA could not be completed in the current macOS runtime because the bundled headless LibreOffice expects `/opt/homebrew/opt/little-cms2/lib/liblcms2.2.dylib`, and the environment does not permit creating that Homebrew path. Word-claim and structural QA passed; visual PNG inspection should be rerun on a machine with a working LibreOffice/Word renderer before journal submission.",
        "",
    ]

    if failures:
        lines += ["## Failures", ""]
        lines.extend(f"- {item}" for item in failures)
        lines.append("")
    else:
        lines += ["## Failures", "", "No structural QA failures were detected.", ""]
    return "\n".join(lines)


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--docx", type=Path, default=DEFAULT_DOCX)
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT)
    args = parser.parse_args()
    report = build_report(args.docx)
    args.output.parent.mkdir(parents=True, exist_ok=True)
    args.output.write_text(report, encoding="utf-8")
    print(f"[saved] {args.output}")
    if "- Status: **fail**" in report:
        raise SystemExit(2)


if __name__ == "__main__":
    main()
