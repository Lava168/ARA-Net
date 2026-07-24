#!/usr/bin/env python3
"""Build the ARA-Net V6 full manuscript DOCX from the public Markdown draft."""
from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


DEFAULT_SOURCE = Path("reports/v6_final_model/manuscript_v6_full_draft.md")
DEFAULT_OUTPUT = Path("reports/v6_final_model/ARA-Net_V6_full_manuscript_draft.docx")

FONT = "Calibri"
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
GRAY = RGBColor(89, 89, 89)
LIGHT_FILL = "F2F4F7"
BORDER = "C9D1D9"


def set_run_font(run, *, size: float | None = None, bold: bool | None = None,
                 italic: bool | None = None, color: RGBColor | None = None) -> None:
    run.font.name = FONT
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if italic is not None:
        run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color


def set_paragraph_spacing(paragraph, *, before: float = 0, after: float = 6,
                          line: float = 1.10) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_borders(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ["top", "left", "bottom", "right"]:
        tag = qn(f"w:{edge}")
        border = borders.find(tag)
        if border is None:
            border = OxmlElement(f"w:{edge}")
            borders.append(border)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), BORDER)


def set_table_width(table, width_dxa: int, col_widths_dxa: list[int]) -> None:
    table.allow_autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(width_dxa))

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), "120")

    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in col_widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = col_widths_dxa[min(idx, len(col_widths_dxa) - 1)]
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(width))


def set_row_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def set_table_cell_margins(table, top: int = 80, bottom: int = 80,
                           start: int = 120, end: int = 120) -> None:
    tbl_pr = table._tbl.tblPr
    margins = tbl_pr.find(qn("w:tblCellMar"))
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    for name, value in [("top", top), ("bottom", bottom), ("start", start), ("end", end)]:
        node = margins.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def strip_inline_md(text: str) -> str:
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"\*(.*?)\*", r"\1", text)
    text = re.sub(r"`([^`]*)`", r"\1", text)
    text = text.replace("\\(", "").replace("\\)", "")
    text = text.replace("\\[", "").replace("\\]", "")
    for old, new in [
        ("\\mathrm{CN}", "CN"),
        ("\\mathrm{MCI}", "MCI"),
        ("\\mathrm{AD}", "AD"),
        ("\\epsilon", "epsilon"),
        ("\\ge", ">="),
        ("\\le", "<="),
        ("\\sum_m", "sum_m"),
        ("\\in", "in"),
        ("\\{", "{"),
        ("\\}", "}"),
    ]:
        text = text.replace(old, new)
    text = re.sub(r"\\([A-Za-z]+)", r"\1", text)
    return text


def readable_formula(expr: str) -> str:
    compact = " ".join(expr.strip().split())
    if compact.startswith("z_{i,k}"):
        return (
            "z_{i,k} = (1/T) sum_{m=1..M} w_m log(max(p_{m,k}(x_i), epsilon)) "
            "+ b_k, for k in {CN, MCI, AD}"
        )
    if "\\bar{p}" in compact and "\\hat{y}" in compact:
        return "p_bar_{s,k} = (1/n_s) sum_{i in s} p_tilde_{i,k}; y_hat_s = argmax_k p_bar_{s,k}"
    if "\\tilde{p}" in compact:
        return "p_tilde_{i,k} = exp(z_{i,k}) / sum_c exp(z_{i,c})"

    text = compact
    for old, new in [
        ("\\tilde{p}", "p_tilde"),
        ("\\bar{p}", "p_bar"),
        ("\\hat{y}", "y_hat"),
        ("\\mathrm{CN}", "CN"),
        ("\\mathrm{MCI}", "MCI"),
        ("\\mathrm{AD}", "AD"),
        ("\\epsilon", "epsilon"),
        ("\\qquad", "    "),
        ("\\left", ""),
        ("\\right", ""),
        ("\\exp", "exp"),
        ("\\log", "log"),
        ("\\max", "max"),
        ("\\sum_c", "sum_c"),
        ("\\arg\\max_k", "argmax_k"),
        ("\\in", "in"),
    ]:
        text = text.replace(old, new)

    frac_re = re.compile(r"\\frac\{([^{}]+)\}\{([^{}]+)\}")
    while True:
        updated = frac_re.sub(r"(\1)/(\2)", text)
        if updated == text:
            break
        text = updated
    text = text.replace("\\sum_{m=1}^{M}", "sum_{m=1..M}")
    text = re.sub(r"\\([A-Za-z]+)", r"\1", text)
    text = text.replace("{", "").replace("}", "")
    return " ".join(text.split())


def add_inline_runs(paragraph, text: str, *, base_size: float = 11.0) -> None:
    pattern = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(strip_inline_md(text[pos:match.start()]))
            set_run_font(run, size=base_size)
        token = match.group(0)
        content = token.strip("*`")
        run = paragraph.add_run(content)
        if token.startswith("**"):
            set_run_font(run, size=base_size, bold=True)
        elif token.startswith("*"):
            set_run_font(run, size=base_size, italic=True)
        else:
            set_run_font(run, size=base_size)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(strip_inline_md(text[pos:]))
        set_run_font(run, size=base_size)


def add_paragraph(doc: Document, text: str, style: str | None = None,
                  *, before: float = 0, after: float = 6,
                  line: float = 1.10, base_size: float = 11.0):
    paragraph = doc.add_paragraph(style=style)
    add_inline_runs(paragraph, text, base_size=base_size)
    set_paragraph_spacing(paragraph, before=before, after=after, line=line)
    return paragraph


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        line = lines[i].strip()
        cells = [strip_inline_md(cell.strip()) for cell in line.strip("|").split("|")]
        if not all(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in cells):
            rows.append(cells)
        i += 1
    return rows, i


def compute_widths(rows: list[list[str]], total: int = 9360) -> list[int]:
    cols = max(len(row) for row in rows)
    header = [cell.lower().replace("\n", " ").strip() for cell in rows[0]]
    header_text = " ".join(header)
    if cols == 10 and "model/protocol" in header_text and "endpoint" in header_text:
        return [1180, 700, 850, 820, 620, 620, 760, 1180, 1040, 1590]
    if cols == 4 and header[:2] == ["split", "subjects/scans"]:
        return [2600, 1500, 1900, 3360]
    if cols == 5 and header and header[0].startswith("true"):
        return [1800, 1900, 1900, 1900, 1860]

    weights = [1] * cols
    for row in rows:
        for idx, cell in enumerate(row):
            weights[idx] = max(weights[idx], min(36, max(4, len(cell) // 8 + 4)))
    weight_sum = sum(weights)
    widths = [max(720, int(total * weight / weight_sum)) for weight in weights]
    delta = total - sum(widths)
    widths[-1] += delta
    return widths


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    cols = max(len(row) for row in rows)
    if cols >= 8:
        doc.add_page_break()
    table = doc.add_table(rows=len(rows), cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    widths = compute_widths(rows)
    set_table_width(table, 9360, widths)
    set_table_cell_margins(table, start=90 if cols >= 8 else 120, end=90 if cols >= 8 else 120)

    for r_idx, row in enumerate(rows):
        set_row_cant_split(table.rows[r_idx])
        if r_idx == 0:
            set_repeat_table_header(table.rows[r_idx])
        for c_idx in range(cols):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            text = row[c_idx] if c_idx < len(row) else ""
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx > 0 and len(text) < 18 else WD_ALIGN_PARAGRAPH.LEFT
            set_paragraph_spacing(paragraph, before=0, after=0, line=1.05)
            add_inline_runs(paragraph, text, base_size=8.4 if cols >= 8 else 9.0)
            set_cell_borders(cell)
            if r_idx == 0:
                set_cell_shading(cell, LIGHT_FILL)
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(31, 58, 95)
    doc.add_paragraph()


def setup_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("ARA-Net V6 manuscript draft - research prototype, not for clinical use")
    set_run_font(footer_run, size=8.5, color=GRAY)


def build_docx(source: Path, output: Path) -> None:
    doc = Document()
    setup_styles(doc)
    lines = source.read_text(encoding="utf-8").splitlines()
    i = 0
    in_math = False
    math_buffer: list[str] = []

    while i < len(lines):
        raw = lines[i]
        text = raw.strip()
        if not text:
            i += 1
            continue

        if text == "\\[":
            in_math = True
            math_buffer = []
            i += 1
            continue
        if in_math:
            if text == "\\]":
                paragraph = add_paragraph(
                    doc,
                    readable_formula(" ".join(math_buffer)),
                    before=3,
                    after=8,
                    base_size=10.5,
                )
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                in_math = False
            else:
                math_buffer.append(text)
            i += 1
            continue

        if text.startswith("|"):
            rows, i = parse_table(lines, i)
            add_table(doc, rows)
            continue

        if text.startswith("# "):
            title = strip_inline_md(text[2:].strip())
            paragraph = doc.add_paragraph()
            set_paragraph_spacing(paragraph, before=0, after=3, line=1.05)
            run = paragraph.add_run(title)
            set_run_font(run, size=22, bold=True, color=DARK_BLUE)
            subtitle = doc.add_paragraph()
            set_paragraph_spacing(subtitle, before=0, after=14, line=1.10)
            run = subtitle.add_run("Full V6 manuscript replacement draft")
            set_run_font(run, size=11, color=GRAY)
        elif text.startswith("## "):
            add_paragraph(doc, strip_inline_md(text[3:].strip()), style="Heading 1")
        elif text.startswith("### "):
            add_paragraph(doc, strip_inline_md(text[4:].strip()), style="Heading 2")
        elif re.match(r"^\d+\.\s+", text):
            item = re.sub(r"^\d+\.\s+", "", text)
            paragraph = add_paragraph(doc, item, style=None, after=4, line=1.10)
            paragraph.style = "List Number"
        elif text.startswith("- "):
            paragraph = add_paragraph(doc, text[2:], style=None, after=4, line=1.10)
            paragraph.style = "List Bullet"
        else:
            add_paragraph(doc, text, after=6, line=1.10)
        i += 1

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output))


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--source", type=Path, default=DEFAULT_SOURCE)
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT)
    args = parser.parse_args()
    build_docx(args.source, args.output)
    print(f"[saved] {args.output}")


if __name__ == "__main__":
    main()
